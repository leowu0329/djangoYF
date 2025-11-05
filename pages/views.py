from django.shortcuts import render, get_object_or_404, redirect
from django.views.generic import TemplateView, ListView, DetailView, CreateView, UpdateView, DeleteView
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.db.models import Q, Case, When, Value, BooleanField, F, Max # Import required database functions
from datetime import date, timedelta
from django.urls import reverse_lazy, reverse
from django.contrib import messages
from django.views.decorators.clickjacking import xframe_options_exempt
from .models import Cases, Land, Build, Person, Survey, FinalDecision, Result, ObjectBuild, Bouns, Auction, City, Township, OfficialDocuments
from users.models import CustomUser # Import CustomUser model
from .forms import CasesForm, LandForm, BuildForm, PersonForm, SurveyForm, FinalDecisionForm, ResultForm, ObjectBuildForm, BounsForm, AuctionForm, OfficialDocumentForm
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_GET
# Import for Word document export
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from io import BytesIO
from urllib.parse import quote
import re

@require_GET
def load_townships(request):
    city_id = request.GET.get('city_id')
    if not city_id:
        return JsonResponse({'error': '缺少城市ID參數'}, status=400)
    try:
        city_id = int(city_id)
    except ValueError:
        return JsonResponse({'error': '無效的城市ID參數'}, status=400)
    try:
        qs = Township.objects.filter(city_id=city_id).order_by('name').values('id', 'name')
        return JsonResponse(list(qs), safe=False)
    except Exception as e:
        return JsonResponse({'error': f'載入鄉鎮資料時發生錯誤: {str(e)}'}, status=500)

@require_GET
def get_city_for_township(request):
    """給定 township_id 回傳該鄉鎮所屬的 city（{id,name}）供前端將 city select 更新。"""
    township_id = request.GET.get('township_id')
    if not township_id:
        return JsonResponse({'error': '缺少鄉鎮ID參數'}, status=400)
    try:
        township_id = int(township_id)
    except ValueError:
        return JsonResponse({'error': '無效的鄉鎮ID參數'}, status=400)
    try:
        t = Township.objects.select_related('city').get(pk=township_id)
    except Township.DoesNotExist:
        return JsonResponse({'error': '找不到指定的鄉鎮資料'}, status=404)
    except Exception as e:
        return JsonResponse({'error': f'載入城市資料時發生錯誤: {str(e)}'}, status=500)
    city = t.city
    return JsonResponse({'id': city.id, 'name': city.name})

class CaseListView(LoginRequiredMixin, ListView):
    model = Cases
    template_name = 'cases/case_list.html'  # <app>/<model>_list.html
    context_object_name = 'cases'
    ordering = ['caseNumber'] # Changed default ordering
    paginate_by = 10  # Default pagination

    def get_paginate_by(self, queryset):
        return self.request.GET.get('page_size', self.paginate_by)

    def get_queryset(self):
        # Annotate latest stopBuyDate for each case
        queryset = Cases.objects.annotate(
            latest_stop_buy_date=Max('results__stopBuyDate')
        ).select_related('city', 'township', 'user').prefetch_related('results', 'finaldecisions')

        # Define the two-week window for stopBuyDate
        today = date.today()
        two_weeks_from_now = today + timedelta(weeks=2)

        # Annotate if the latest_stop_buy_date is within the future two-week window
        queryset = queryset.annotate(
            is_future_two_weeks_stop_buy=Case(
                When(
                    latest_stop_buy_date__isnull=False,
                    latest_stop_buy_date__gte=today,
                    latest_stop_buy_date__lte=two_weeks_from_now,
                    then=Value(True)
                ),
                default=Value(False),
                output_field=BooleanField()
            )
        )

        # Filtering
        city_filter = self.request.GET.get('city')
        user_filter = self.request.GET.get('user')

        if city_filter:
            queryset = queryset.filter(city__id=city_filter)
        if user_filter:
            queryset = queryset.filter(user__id=user_filter)

        # Keyword Search
        search_query = self.request.GET.get('q')
        if search_query:
            queryset = queryset.filter(
                Q(caseNumber__icontains=search_query) |
                Q(city__name__icontains=search_query) |
                Q(township__name__icontains=search_query) |
                Q(village__icontains=search_query) |
                Q(neighbor__icontains=search_query) |
                Q(street__icontains=search_query) |
                Q(section__icontains=search_query) |
                Q(lane__icontains=search_query) |
                Q(alley__icontains=search_query) |
                Q(number__icontains=search_query) |
                Q(Floor__icontains=search_query) |
                Q(user__username__icontains=search_query) |
                Q(results__stopBuyDate__icontains=search_query) | 
                Q(finaldecisions__finalDecision__icontains=search_query) |
                Q(results__actionResult__icontains=search_query) |
                Q(results__objectNumber__icontains=search_query) |
                Q(officialdocuments__docNumber__icontains=search_query)
            ).distinct()

        # Sorting logic
        sort_by = self.request.GET.get('sort_by')
        order = self.request.GET.get('order', 'desc')

        ordering_criteria = []

        # Primary sorting: Cases with future two-week stopBuyDate at the top, sorted by stopBuyDate DESC
        # Then other cases sorted by updated DESC

        if sort_by == 'stopBuyDate':
            # If sorting by stopBuyDate is explicitly requested
            if order == 'asc':
                ordering_criteria.append(F('is_future_two_weeks_stop_buy').desc()) # True first
                ordering_criteria.append(F('latest_stop_buy_date').asc(nulls_last=True))
            else: # Default to desc for stopBuyDate
                ordering_criteria.append(F('is_future_two_weeks_stop_buy').desc()) # True first
                ordering_criteria.append(F('latest_stop_buy_date').desc(nulls_last=True))
        elif sort_by == 'caseNumber':
            # If sorting by caseNumber is requested, prioritize stopBuyDate then caseNumber
            ordering_criteria.append(F('is_future_two_weeks_stop_buy').desc()) # True first
            ordering_criteria.append(F('latest_stop_buy_date').desc(nulls_last=True)) # Then stopBuyDate DESC
            if order == 'asc':
                ordering_criteria.append(F('caseNumber').asc())
            else:
                ordering_criteria.append(F('caseNumber').desc())
        elif sort_by == 'updated':
            # If sorting by updated date is requested, prioritize stopBuyDate then updated
            ordering_criteria.append(F('is_future_two_weeks_stop_buy').desc()) # True first
            ordering_criteria.append(F('latest_stop_buy_date').desc(nulls_last=True)) # Then stopBuyDate DESC
            if order == 'asc':
                ordering_criteria.append(F('updated').asc())
            else:
                ordering_criteria.append(F('updated').desc())
        else:
            # Default sorting: prioritize stopBuyDate then updated date
            ordering_criteria.append(F('is_future_two_weeks_stop_buy').desc()) # True first
            ordering_criteria.append(F('latest_stop_buy_date').desc(nulls_last=True)) # Then stopBuyDate DESC
            ordering_criteria.append(F('updated').desc()) # Then updated DESC for other cases

        queryset = queryset.order_by(*ordering_criteria).distinct()

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        today = date.today()
        two_weeks_from_now = today + timedelta(weeks=2)
        
        # Pass sorting parameters to context for template to use
        context['sort_by'] = self.request.GET.get('sort_by', 'stopBuyDate') # Default to stopBuyDate for initial view
        context['order'] = self.request.GET.get('order', 'desc')
        context['search_query'] = self.request.GET.get('q', '')
        context['selected_city'] = self.request.GET.get('city', '')
        context['selected_user'] = self.request.GET.get('user', '')
        context['page_size'] = self.request.GET.get('page_size', self.paginate_by)

        for case in context['cases']:
            latest_stop_buy_date = getattr(case, 'latest_stop_buy_date', None)
            if latest_stop_buy_date and today <= latest_stop_buy_date <= two_weeks_from_now:
                case.display_stop_buy_date = latest_stop_buy_date.strftime("%Y-%m-%d")
                case.stop_buy_date_style = "future-two-weeks"
            else:
                case.display_stop_buy_date = "—"
                case.stop_buy_date_style = ""
        return context

class CaseDetailView(LoginRequiredMixin, DetailView):
    model = Cases
    template_name = 'cases/case_detail.html'
    context_object_name = 'case'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        case = self.get_object()
        context['officialdocuments'] = case.officialdocuments.all()
        return context

class CaseCreateView(LoginRequiredMixin, CreateView):
    model = Cases
    template_name = 'cases/case_form.html'
    form_class = CasesForm

    @xframe_options_exempt
    def dispatch(self, *args, **kwargs):
        return super().dispatch(*args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if self.request.GET.get('is_iframe'):
            context['is_modal_form'] = True
        return context

    def form_valid(self, form):
        form.instance.user = self.request.user
        messages.success(self.request, f"案件 '{form.instance.caseNumber}' 建立成功！")
        response = super().form_valid(form)
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest' or self.request.GET.get('is_iframe'):
            return render(self.request, 'cases/case_form_success.html', {'case_id': form.instance.id})
        return response

    def form_invalid(self, form):
        messages.error(self.request, "建立案件失敗，請檢查輸入。")
        return super().form_invalid(form)

class CaseUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = Cases
    template_name = 'cases/case_form.html'
    form_class = CasesForm

    @xframe_options_exempt
    def dispatch(self, *args, **kwargs):
        return super().dispatch(*args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if self.request.GET.get('is_iframe'):
            context['is_modal_form'] = True
        return context

    def form_valid(self, form):
        # form.instance.user = self.request.user # 移除或註解此行以避免更新負責人
        messages.success(self.request, f"案件 '{form.instance.caseNumber}' 更新成功！")
        response = super().form_valid(form)
        if self.request.headers.get('x-requested-with') == 'XMLHttpRequest' or self.request.GET.get('is_iframe'):
            return render(self.request, 'cases/case_form_success.html', {'case_id': form.instance.id})
        return response

    def test_func(self):
        case = self.get_object()
        if self.request.user == case.user or self.request.user.is_staff:
            return True
        return False

class CaseDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    model = Cases
    template_name = 'cases/case_confirm_delete.html'
    context_object_name = 'case'
    success_url = reverse_lazy('case_list')

    def test_func(self):
        case = self.get_object()
        if self.request.user == case.user or self.request.user.is_staff:
            return True
        return False

    def delete(self, request, *args, **kwargs):
        case = self.get_object()
        messages.success(self.request, f"案件 '{case.caseNumber}' 已成功刪除！")
        return super().delete(request, *args, **kwargs)

# Land CRUD
class LandCreateView(LoginRequiredMixin, CreateView):
    model = Land
    form_class = LandForm
    template_name = 'lands/land_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.case = get_object_or_404(Cases, pk=kwargs.get('case_pk'))
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['case'] = self.case
        return context

    def form_valid(self, form):
        form.instance.cases = self.case
        messages.success(self.request, "土地資料已新增！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case.pk])

class LandUpdateView(LoginRequiredMixin, UpdateView):
    model = Land
    form_class = LandForm
    template_name = 'lands/land_form.html'

    def form_valid(self, form):
        messages.success(self.request, "土地資料已更新！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.object.cases_id])

class LandDeleteView(LoginRequiredMixin, DeleteView):
    model = Land
    template_name = 'lands/land_confirm_delete.html'

    def get_success_url(self):
        return reverse('case_detail', args=[self.object.cases.pk])

    def delete(self, request, *args, **kwargs):
        self.object = self.get_object()
        case_id = self.object.cases_id
        messages.success(self.request, "土地資料已刪除！")
        return super().delete(request, *args, **kwargs)

# Build CRUD
class BuildCreateView(LoginRequiredMixin, CreateView):
    model = Build
    form_class = BuildForm
    template_name = 'builds/build_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.case = get_object_or_404(Cases, pk=kwargs.get('case_pk'))
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['case'] = self.case
        return context

    def form_valid(self, form):
        form.instance.cases = self.case
        messages.success(self.request, "建物資料已新增！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case.pk])

class BuildUpdateView(LoginRequiredMixin, UpdateView):
    model = Build
    form_class = BuildForm
    template_name = 'builds/build_form.html'

    def form_valid(self, form):
        messages.success(self.request, "建物資料已更新！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.object.cases_id])

class BuildDeleteView(LoginRequiredMixin, DeleteView):
    model = Build
    template_name = 'builds/build_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        self.case_pk = self.object.cases.pk
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case_pk])

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, "建物資料已刪除！")
        return super().delete(request, *args, **kwargs)

# Person CRUD
class PersonCreateView(LoginRequiredMixin, CreateView):
    model = Person
    form_class = PersonForm
    template_name = 'people/person_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.case = get_object_or_404(Cases, pk=kwargs.get('case_pk'))
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['case'] = self.case
        return context

    def form_valid(self, form):
        form.instance.cases = self.case
        messages.success(self.request, "人員資料已新增！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case.pk])

class PersonUpdateView(LoginRequiredMixin, UpdateView):
    model = Person
    form_class = PersonForm
    template_name = 'people/person_form.html'

    def form_valid(self, form):
        messages.success(self.request, "人員資料已更新！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.object.cases_id])

class PersonDeleteView(LoginRequiredMixin, DeleteView):
    model = Person
    template_name = 'people/person_confirm_delete.html'

    def get_success_url(self):
        return reverse('case_detail', args=[self.object.cases.pk])

    def delete(self, request, *args, **kwargs):
        self.object = self.get_object()
        case_id = self.object.cases_id
        messages.success(self.request, "人員資料已刪除！")
        return super().delete(request, *args, **kwargs)

# Survey CRUD
class SurveyCreateView(LoginRequiredMixin, CreateView):
    model = Survey
    form_class = SurveyForm
    template_name = 'surveys/survey_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.case = get_object_or_404(Cases, pk=kwargs.get('case_pk'))
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['case'] = self.case
        return context

    def form_valid(self, form):
        form.instance.cases = self.case
        messages.success(self.request, "勘查資料已新增！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case.pk])

class SurveyUpdateView(LoginRequiredMixin, UpdateView):
    model = Survey
    form_class = SurveyForm
    template_name = 'surveys/survey_form.html'

    def form_valid(self, form):
        messages.success(self.request, "勘查資料已更新！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.object.cases_id])

class SurveyDeleteView(LoginRequiredMixin, DeleteView):
    model = Survey
    template_name = 'surveys/survey_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        self.case_pk = self.object.cases.pk
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case_pk])

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, "勘查資料已刪除！")
        return super().delete(request, *args, **kwargs)

# FinalDecision CRUD
class FinalDecisionCreateView(LoginRequiredMixin, CreateView):
    model = FinalDecision
    form_class = FinalDecisionForm
    template_name = 'finaldecisions/finaldecision_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.case = get_object_or_404(Cases, pk=kwargs.get('case_pk'))
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['case'] = self.case
        return context

    def form_valid(self, form):
        form.instance.cases = self.case
        messages.success(self.request, "最終判定已新增！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case.pk])

class FinalDecisionUpdateView(LoginRequiredMixin, UpdateView):
    model = FinalDecision
    form_class = FinalDecisionForm
    template_name = 'finaldecisions/finaldecision_form.html'

    def form_valid(self, form):
        messages.success(self.request, "最終判定已更新！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.object.cases_id])

class FinalDecisionDeleteView(LoginRequiredMixin, DeleteView):
    model = FinalDecision
    template_name = 'finaldecisions/finaldecision_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        self.case_pk = self.object.cases.pk
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case_pk])

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, "最終判定已刪除！")
        return super().delete(request, *args, **kwargs)

# Result CRUD
class ResultCreateView(LoginRequiredMixin, CreateView):
    model = Result
    form_class = ResultForm
    template_name = 'results/result_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.case = get_object_or_404(Cases, pk=kwargs.get('case_pk'))
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['case'] = self.case
        return context

    def form_valid(self, form):
        form.instance.cases = self.case
        messages.success(self.request, "結果資料已新增！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case.pk])

class ResultUpdateView(LoginRequiredMixin, UpdateView):
    model = Result
    form_class = ResultForm
    template_name = 'results/result_form.html'

    def form_valid(self, form):
        messages.success(self.request, "結果資料已更新！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.object.cases_id])

class ResultDeleteView(LoginRequiredMixin, DeleteView):
    model = Result
    template_name = 'results/result_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        self.case_pk = self.object.cases.pk
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case_pk])

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, "結果資料已刪除！")
        return super().delete(request, *args, **kwargs)

# ObjectBuild CRUD
class ObjectBuildCreateView(LoginRequiredMixin, CreateView):
    model = ObjectBuild
    form_class = ObjectBuildForm
    template_name = 'objectbuilds/objectbuild_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.case = get_object_or_404(Cases, pk=kwargs.get('case_pk'))
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['case'] = self.case
        return context

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['initial'] = {'cases': self.case.id}
        return kwargs

    def form_valid(self, form):
        form.instance.cases = self.case
        try:
            messages.success(self.request, "比價建物已新增！")
            response = super().form_valid(form)
            return response
        except Exception as e:
            print(f"ObjectBuildCreateView: Error saving ObjectBuild: {e}")
            messages.error(self.request, "儲存比價建物資料時發生錯誤。")
            return self.form_invalid(form)

    def form_invalid(self, form):
        print("ObjectBuildCreateView: form_invalid - Errors:", form.errors) # Debugging line
        messages.error(self.request, "比價建物表單驗證失敗，請檢查輸入。")
        return super().form_invalid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case.pk])

class ObjectBuildUpdateView(LoginRequiredMixin, UpdateView):
    model = ObjectBuild
    form_class = ObjectBuildForm
    template_name = 'objectbuilds/objectbuild_form.html'

    def form_valid(self, form):
        print("ObjectBuildUpdateView: form_valid - Before save")
        try:
            messages.success(self.request, "比價建物已更新！")
            response = super().form_valid(form)
            print("ObjectBuildUpdateView: form_valid - After save")
            return response
        except Exception as e:
            print(f"ObjectBuildUpdateView: Error saving ObjectBuild: {e}")
            messages.error(self.request, "儲存比價建物資料時發生錯誤。")
            return self.form_invalid(form)

    def form_invalid(self, form):
        print("ObjectBuildUpdateView: form_invalid - Errors:", form.errors) # Debugging line
        messages.error(self.request, "比價建物表單驗證失敗，請檢查輸入。")
        return super().form_invalid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.object.cases_id])

class ObjectBuildDeleteView(LoginRequiredMixin, DeleteView):
    model = ObjectBuild
    template_name = 'objectbuilds/objectbuild_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        self.case_pk = self.object.cases.pk
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case_pk])

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, "比價建物已刪除！")
        return super().delete(request, *args, **kwargs)

# Bouns CRUD
class BounsCreateView(LoginRequiredMixin, CreateView):
    model = Bouns
    form_class = BounsForm
    template_name = 'bouns/bouns_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.objectbuild = get_object_or_404(ObjectBuild, pk=kwargs.get('objectbuild_pk'))
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['objectbuild'] = self.objectbuild
        return context

    def get_initial(self):
        initial = super().get_initial()
        if self.request.user.is_authenticated:
            initial['bounsPerson'] = self.request.user
        return initial

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        initial = self.get_initial()  # Get initial data including bounsPerson
        initial['objectbuild'] = self.objectbuild.id  # Add objectbuild to initial data
        kwargs['initial'] = initial  # Set the combined initial data
        return kwargs

    def form_valid(self, form):
        print("form_valid: Before assigning objectbuild")
        try:
            form.instance.objectbuild = self.objectbuild
            print(f"form_valid: After assigning objectbuild. objectbuild ID: {form.instance.objectbuild.id}")
            # messages.success(self.request, "加成資料已新增！") # Temporarily removed
            response = super().form_valid(form)
            print("form_valid: After super().form_valid")
            return response
        except Exception as e:
            print(f"Error saving Bouns: {e}")
            messages.error(self.request, "儲存加成資料時發生錯誤。")
            return self.form_invalid(form)

    def form_invalid(self, form):
        print(form.errors) # Debugging line to show form errors
        messages.error(self.request, "表單驗證失敗，請檢查輸入。")
        return super().form_invalid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.objectbuild.cases_id])

class BounsUpdateView(LoginRequiredMixin, UpdateView):
    model = Bouns
    form_class = BounsForm
    template_name = 'bouns/bouns_form.html'

    def form_valid(self, form):
        messages.success(self.request, "加成資料已更新！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.object.objectbuild.cases_id])

class BounsDeleteView(LoginRequiredMixin, DeleteView):
    model = Bouns
    template_name = 'bouns/bouns_confirm_delete.html'

    def delete(self, request, *args, **kwargs):
        self.object = self.get_object()
        case_id = self.object.objectbuild.cases_id
        messages.success(self.request, "加成資料已刪除！")
        response = super().delete(request, *args, **kwargs)
        self.success_url = reverse_lazy('case_detail', args=[case_id])
        return response

# Auction CRUD
class AuctionCreateView(LoginRequiredMixin, CreateView):
    model = Auction
    form_class = AuctionForm
    template_name = 'auctions/auction_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.case = get_object_or_404(Cases, pk=kwargs.get('case_pk'))
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['case'] = self.case
        return context

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['initial'] = {'cases': self.case.id}
        return kwargs

    def form_valid(self, form):
        form.instance.cases = self.case
        messages.success(self.request, "拍賣資料已新增！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case.pk])

class AuctionUpdateView(LoginRequiredMixin, UpdateView):
    model = Auction
    form_class = AuctionForm
    template_name = 'auctions/auction_form.html'

    def form_valid(self, form):
        messages.success(self.request, "拍賣資料已更新！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.object.cases_id])

class AuctionDeleteView(LoginRequiredMixin, DeleteView):
    model = Auction
    template_name = 'auctions/auction_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        self.case_pk = self.object.cases.pk
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case_pk])

    def delete(self, request, *args, **kwargs):
        messages.success(self.request, "拍賣資料已刪除！")
        return super().delete(request, *args, **kwargs)

# OfficialDocument CRUD
class OfficialDocumentCreateView(LoginRequiredMixin, CreateView):
    model = OfficialDocuments
    form_class = OfficialDocumentForm
    template_name = 'OfficialDocuments/officialdocument_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.case = get_object_or_404(Cases, pk=kwargs.get('case_pk'))
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['case'] = self.case
        return context

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['initial'] = {'cases': self.case.id}
        return kwargs

    def form_valid(self, form):
        form.instance.cases = self.case
        messages.success(self.request, "公文資料已新增！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case.pk])

class OfficialDocumentDetailView(LoginRequiredMixin, DetailView):
    model = OfficialDocuments
    template_name = 'OfficialDocuments/officialdocument_detail.html'
    context_object_name = 'officialdocument'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        return context

class OfficialDocumentUpdateView(LoginRequiredMixin, UpdateView):
    model = OfficialDocuments
    form_class = OfficialDocumentForm
    template_name = 'OfficialDocuments/officialdocument_form.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['case'] = self.object.cases
        return context

    def form_valid(self, form):
        messages.success(self.request, "公文資料已更新！")
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('case_detail', args=[self.object.cases.pk])

class OfficialDocumentDeleteView(LoginRequiredMixin, DeleteView):
    model = OfficialDocuments
    template_name = 'OfficialDocuments/officialdocument_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        self.case_pk = self.object.cases.pk
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse('case_detail', args=[self.case_pk])


def set_cell_shading(cell, color):
    """
    Set background shading for a docx table cell.
    :param cell: The cell to shade.
    :param color: Hex color string, e.g., 'FFDAB9' (no #).
    """
    from docx.oxml import OxmlElement
    
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Find existing shd element
    shd = tcPr.find(qn('w:shd'))
    
    # Create new shd element if it doesn't exist
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    
    # Set shading properties
    shd.set(qn('w:fill'), color)
    shd.set(qn('w:color'), "auto")
    shd.set(qn('w:val'), "clear")

def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph in a python-docx document.
    :param paragraph: The paragraph to add the hyperlink to.
    :param text: The display text for the hyperlink.
    :param url: The URL for the hyperlink.
    """
    # This function uses the lxml library to add a hyperlink to a docx paragraph.
    # python-docx does not support hyperlinks natively.
    import docx
    from docx.oxml import OxmlElement

    # Create the w:hyperlink tag and set the r:id attribute
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a w:rPr element
    rPr = OxmlElement('w:rPr')

    # Add formatting (blue color and underline)
    # Color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    # Underline
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    new_run.append(rPr)

    # Create a w:t element
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def case_export_word(request, pk):
    case = get_object_or_404(Cases, pk=pk)

    document = Document()
    # Set page orientation to landscape
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    # Set page margins to 1.27 cm on all sides
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)

    # Set default fonts for Chinese, English and numbers
    style = document.styles['Normal']
    style.font.name = 'Arial' # Set ASCII font for English and numbers
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體') # Set East Asia font for Chinese

    document.add_heading(f"案件報告 - {case.caseNumber}", level=1)

    # Debugging: Print the case number to see its value
    print(f"DEBUG: case.caseNumber before filename sanitization: {case.caseNumber}")

    # Sanitize the case number for use in filename
    cleaned_case_number = re.sub(r'[\\/:*?"<>|]', '_', case.caseNumber)
    base_filename = f"{cleaned_case_number}.docx"

    # For filename parameter (legacy support, typically needs to be ASCII or Latin-1 compatible)
    # Some browsers might not like non-ASCII in 'filename', even if 'filename*' is present.
    # We will try to encode it as Latin-1. If it fails, fallback to something else.
    try:
        filename_ascii = base_filename.encode('latin-1').decode('latin-1')
    except UnicodeEncodeError:
        # Fallback if latin-1 encoding is not possible (e.g., many non-ASCII chars)
        # This fallback might still cause issues on some very old browsers, but with filename* it should be fine.
        filename_ascii = re.sub(r'[^\x00-\x7F]+', '_', base_filename) # Replace non-ASCII with underscore

    # For filename* parameter (RFC 6266, handles UTF-8 characters correctly)
    encoded_filename_utf8 = quote(base_filename, encoding='utf-8')

    # 1. 基本資料
    document.add_heading("基本資料", level=2)
    table_basic = document.add_table(rows=2, cols=5) # Changed to 5 columns
    table_basic.autofit = False # Set autofit to False to manually control column widths
    table_basic.style = 'Table Grid' # Add this line to apply a table style with borders
    # Removed: table_basic.columns[0].width = Inches(0.5)

    # Calculate available content width (assuming A4 landscape, 1.27cm margins)
    # A4 landscape width is 29.7 cm, height is 21 cm
    # Usable width = 29.7 cm - (2 * 1.27 cm margin) = 29.7 - 2.54 = 27.16 cm
    usable_width_cm = 27.16
    table_basic.width = Cm(usable_width_cm) # Set explicit total table width

    # Define specific column widths in Cm
    specific_column_widths_cm = [
        1.44,  # 狀態 1.44cm
        8.5,   # 案號 8.5cm
        9.5,   # 地址 9.5cm
        4.0,   # 公司 4.0cm
        2.25   # 負責人 2.25cm
    ]

    # Debugging: Print specific column widths
    print(f"DEBUG: Specific column widths (cm): {specific_column_widths_cm}")

    # Set column widths for each cell (column-wise iteration)
    for col_idx, width_cm in enumerate(specific_column_widths_cm):
        for cell in table_basic.columns[col_idx].cells:
            cell.width = Cm(width_cm)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER # Set vertical alignment

    # Row 1 Headers
    hdr_cells_basic_r1 = table_basic.rows[0].cells
    hdr_cells_basic_r1[0].text = "狀態"
    hdr_cells_basic_r1[1].text = "案號"
    hdr_cells_basic_r1[2].text = "地址"
    hdr_cells_basic_r1[3].text = "公司"
    hdr_cells_basic_r1[4].text = "負責人"

    # Make headers bold
    for cell in hdr_cells_basic_r1:
        if cell.text:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    # Row 2 Data
    data_cells_basic_r2 = table_basic.rows[1].cells
    data_cells_basic_r2[0].text = case.status or '—'
    data_cells_basic_r2[1].text = case.caseNumber or '—'
    data_cells_basic_r2[2].text = case.fullAddress() or '—'
    data_cells_basic_r2[3].text = case.company or '—'
    data_cells_basic_r2[4].text = case.user.profile.nickname or case.user.username or '—'

    # Apply shading to header row (Row 0) - Light Orange (FFDAB9)
    for cell in table_basic.rows[0].cells:
        set_cell_shading(cell, 'FFDAB9')

    # Apply shading to data row (Row 1) - White (FFFFFF)
    for cell in table_basic.rows[1].cells:
        set_cell_shading(cell, 'FFFFFF')

    # 2. 人員資訊
    document.add_heading("人員資訊", level=2)
    if case.people.exists():
        table_person = document.add_table(rows=1, cols=3) # Start with 1 row for headers
        table_person.autofit = True
        table_person.style = 'Table Grid'

        # Set header row
        hdr_cells_person = table_person.rows[0].cells
        hdr_cells_person[0].text = "姓名"
        hdr_cells_person[1].text = "分類"
        hdr_cells_person[2].text = "電話"

        # Make headers bold and set background to light orange
        for cell in hdr_cells_person:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            set_cell_shading(cell, 'FFDAB9')

        # Add data rows and set background to white
        for person in case.people.all():
            row_cells = table_person.add_row().cells
            row_cells[0].text = person.name or '—'
            row_cells[1].text = person.type or '—'
            row_cells[2].text = person.phone or '—'
            for cell in row_cells:
                set_cell_shading(cell, 'FFFFFF')
    else:
        document.add_paragraph("無人員記錄。")


    # 3. 土地資訊
    document.add_heading("土地資訊", level=2)
    if case.lands.exists():
        table_land = document.add_table(rows=1, cols=5)
        table_land.autofit = True
        table_land.style = 'Table Grid' # Add this line
        hdr_cells_land = table_land.rows[0].cells
        hdr_cells_land[0].text = "地號"
        hdr_cells_land[1].text = "地坪(平方公尺)"
        hdr_cells_land[2].text = "個人持分"
        hdr_cells_land[3].text = "所有持分"
        hdr_cells_land[4].text = "計算後地坪(坪)"
        # Make headers bold and set background to light orange
        for cell in hdr_cells_land:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            set_cell_shading(cell, 'FFDAB9')

        # Define specific column widths for Land table
        specific_column_widths_cm_land = [
            8.44, # 地號
            4.24, # 地坪(平方公尺)
            4.24, # 個人持分
            4.24, # 所有持分
            4.24, # 計算後地坪(坪)
        ]

        # Set column widths for each cell (column-wise iteration)
        for col_idx, width_cm in enumerate(specific_column_widths_cm_land):
            for cell in table_land.columns[col_idx].cells:
                cell.width = Cm(width_cm)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for land in case.lands.all():
            row_cells = table_land.add_row().cells
            row_cells[0].text = land.formatted_landNumber or '—'
            row_cells[1].text = str(land.area) or '—'
            row_cells[2].text = str(land.holdingPointPersonal) or '—'
            row_cells[3].text = str(land.holdingPointAll) or '—'
            row_cells[4].text = str(land.calculatedArea) or '—'
            # Set data row white background
            for cell in row_cells:
                set_cell_shading(cell, 'FFFFFF')
    else:
        document.add_paragraph("無土地記錄。")


    # 4. 建物資訊
    document.add_heading("建物資訊", level=2)
    if case.builds.exists():
        table_build = document.add_table(rows=1, cols=7)
        table_build.autofit = True
        table_build.style = 'Table Grid' # Add this line
        hdr_cells_build = table_build.rows[0].cells
        hdr_cells_build[0].text = "建號"
        hdr_cells_build[1].text = "建坪(平方公尺)"
        hdr_cells_build[2].text = "個人持分"
        hdr_cells_build[3].text = "所有持分"
        hdr_cells_build[4].text = "計算後建坪"
        hdr_cells_build[5].text = "建物型"
        hdr_cells_build[6].text = "使用分區"
        # Make headers bold and set background to light orange
        for cell in hdr_cells_build:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            set_cell_shading(cell, 'FFDAB9')

        # Define specific column widths for Build table
        specific_column_widths_cm_build = [
            6.19, # 建號
            3.25, # 建坪(平方公尺)
            2.50, # 個人持分
            2.57, # 所有持分
            2.68, # 計算後建坪
            4.58, # 建物型
            3.63, # 使用分區
        ]

        # Set column widths for each cell (column-wise iteration)
        for col_idx, width_cm in enumerate(specific_column_widths_cm_build):
            for cell in table_build.columns[col_idx].cells:
                cell.width = Cm(width_cm)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for build in case.builds.all():
            row_cells = table_build.add_row().cells
            row_cells[0].text = build.formatted_buildNumber or '—'
            row_cells[1].text = str(build.area) or '—'
            row_cells[2].text = str(build.holdingPointPersonal) or '—'
            row_cells[3].text = str(build.holdingPointAll) or '—'
            row_cells[4].text = str(build.calculatedArea) or '—'
            row_cells[5].text = build.typeUse or '—'
            row_cells[6].text = build.usePartition or '—'
            # Set data row white background
            for cell in row_cells:
                set_cell_shading(cell, 'FFFFFF')
    else:
        document.add_paragraph("無建物記錄。")


    # 5. 拍賣資訊
    document.add_heading("拍賣資訊", level=2)
    if case.auctions.exists():
        table_auction = document.add_table(rows=1, cols=9)
        table_auction.autofit = True
        table_auction.style = 'Table Grid' # Add this line
        hdr_cells_auction = table_auction.rows[0].cells
        hdr_cells_auction[0].text = "拍別"
        hdr_cells_auction[1].text = "拍賣日"
        hdr_cells_auction[2].text = "底價"
        hdr_cells_auction[3].text = "坪價"
        hdr_cells_auction[4].text = "CP"
        hdr_cells_auction[5].text = "點閱"
        hdr_cells_auction[6].text = "監控"
        hdr_cells_auction[7].text = "成交件數"
        hdr_cells_auction[8].text = "保証金"
        # Make headers bold and set background to light orange
        for cell in hdr_cells_auction:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            set_cell_shading(cell, 'FFDAB9')

        # Define specific column widths for Auction table
        specific_column_widths_cm_auction = [
            1.69, # 拍別
            5.00, # 拍賣日
            3.50, # 底價
            3.25, # 坪價
            4.25, # CP
            1.75, # 點閱
            1.75, # 監控
            2.00, # 成交件數
            2.21, # 保証金
        ]

        # Set column widths for each cell (column-wise iteration)
        for col_idx, width_cm in enumerate(specific_column_widths_cm_auction):
            for cell in table_auction.columns[col_idx].cells:
                cell.width = Cm(width_cm)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for auction in case.auctions.all():
            row_cells = table_auction.add_row().cells
            row_cells[0].text = auction.type or '—'
            row_cells[1].text = auction.auctionDate.strftime("%Y-%m-%d") if auction.auctionDate else '—'
            row_cells[2].text = str(auction.floorPrice) or '—'
            row_cells[3].text = str(auction.pingPrice) or '—'
            # For CP, include the suggested text based on value
            cp_value_text = str(auction.calculated_cp_value)
            if auction.type == '1拍' and auction.calculated_cp_value > 0.92 or \
               auction.type == '2拍' and auction.calculated_cp_value > 1.15 or \
               auction.type == '3拍' and auction.calculated_cp_value > 1.44 or \
               auction.type == '4拍' and auction.calculated_cp_value > 1.80:
                cp_value_text += " (建議進場)"
            else:
                cp_value_text += " (不可進場)"
            row_cells[4].text = cp_value_text
            row_cells[5].text = str(auction.click) or '—'
            row_cells[6].text = str(auction.monitor) or '—'
            row_cells[7].text = str(auction.caseCount) or '—'
            row_cells[8].text = str(auction.margin) or '—'
            # Set data row white background
            for cell in row_cells:
                set_cell_shading(cell, 'FFFFFF')
    else:
        document.add_paragraph("無拍賣記錄。")


    # 6. 勘查記錄
    document.add_heading("勘查記錄", level=2)
    if case.surveys.exists():
        table_survey = document.add_table(rows=1, cols=10)
        table_survey.autofit = True
        table_survey.style = 'Table Grid' # Add this line
        hdr_cells_survey = table_survey.rows[0].cells
        hdr_cells_survey[0].text = "初勘日"
        hdr_cells_survey[1].text = "會勘日"
        hdr_cells_survey[2].text = "法拍公告"
        hdr_cells_survey[3].text = "998"
        hdr_cells_survey[4].text = "物件照片"
        hdr_cells_survey[5].text = "市場行情"
        hdr_cells_survey[6].text = "法拍記錄"
        hdr_cells_survey[7].text = "現場勘查"
        hdr_cells_survey[8].text = "收發文薄"
        hdr_cells_survey[9].text = "流水帳"
        # Make headers bold and set background to light orange
        for cell in hdr_cells_survey:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            set_cell_shading(cell, 'FFDAB9')

        # Define specific column widths for Survey table
        specific_column_widths_cm_survey = [
            2.84, # 初勘日
            2.85, # 會勘日
            2.46, # 法拍公告
            2.46, # 998
            2.46, # 物件照片
            2.46, # 市場行情
            2.46, # 法拍記錄
            2.46, # 現場勘查
            2.46, # 收發文薄
            2.46, # 流水帳
        ]

        # Set column widths for each cell (column-wise iteration)
        for col_idx, width_cm in enumerate(specific_column_widths_cm_survey):
            for cell in table_survey.columns[col_idx].cells:
                cell.width = Cm(width_cm)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for survey in case.surveys.all():
            row_cells = table_survey.add_row().cells
            row_cells[0].text = survey.firstDay.strftime("%Y-%m-%d") if survey.firstDay else '—'
            row_cells[1].text = survey.secondDay.strftime("%Y-%m-%d") if survey.secondDay else '—'
            # Handle hyperlinks for specific cells
            url_fields = {
                2: survey.foreclosureAnnouncementLink, # 法拍公告
                3: survey.house988Link,               # 998
                4: survey.objectPhotoLink,            # 物件照片
                5: survey.netMarketPriceLink,         # 市場行情
                6: survey.foreclosureRecordLink,      # 法拍記錄
                7: survey.objectViewLink,             # 現場勘查
                8: survey.pagesViewLink,              # 收發文薄
                9: survey.moneytViewLink,             # 流水帳
            }

            for col_idx, url in url_fields.items():
                if url and url != '—':
                    # Clear existing content and add hyperlink
                    row_cells[col_idx].text = '' # Clear existing text
                    add_hyperlink(row_cells[col_idx].paragraphs[0], "連結", url)
                else:
                    row_cells[col_idx].text = '—' # No URL, display as '—'

            # Remaining non-URL fields (if any, but already covered above)
            # row_cells[0].text = survey.firstDay.strftime("%Y-%m-%d") if survey.firstDay else '—'
            # row_cells[1].text = survey.secondDay.strftime("%Y-%m-%d") if survey.secondDay else '—'
            # Set data row white background
            for cell in row_cells:
                set_cell_shading(cell, 'FFFFFF')
    else:
        document.add_paragraph("無勘查記錄。")


    # 7. 最終判定
    document.add_heading("最終判定", level=2)
    if case.finaldecisions.exists():
        table_fd = document.add_table(rows=1, cols=6)
        table_fd.autofit = True
        table_fd.style = 'Table Grid' # Add this line
        hdr_cells_fd = table_fd.rows[0].cells
        hdr_cells_fd[0].text = "最終判定"
        hdr_cells_fd[1].text = "分類"
        hdr_cells_fd[2].text = "人員"
        hdr_cells_fd[3].text = "日期"
        hdr_cells_fd[4].text = "工作轄區"
        hdr_cells_fd[5].text = "備註"
        # Make headers bold and set background to light orange
        for cell in hdr_cells_fd:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            set_cell_shading(cell, 'FFDAB9')

        # Define specific column widths for FinalDecision table
        specific_column_widths_cm_fd = [
            2.44,  # 最終判定
            2.75,  # 分類
            2.50,  # 人員
            3.00,  # 日期
            3.00,  # 工作轄區
            11.71, # 備註
        ]

        # Set column widths for each cell (column-wise iteration)
        for col_idx, width_cm in enumerate(specific_column_widths_cm_fd):
            for cell in table_fd.columns[col_idx].cells:
                cell.width = Cm(width_cm)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for fd in case.finaldecisions.all():
            row_cells = table_fd.add_row().cells
            row_cells[0].text = fd.finalDecision or '—'
            row_cells[1].text = fd.type or '—'
            row_cells[2].text = fd.name or '—'
            row_cells[3].text = fd.date.strftime("%Y-%m-%d") if fd.date else '—'
            row_cells[4].text = fd.workArea or '—'
            row_cells[5].text = fd.remark or '—'
            # Set data row white background
            for cell in row_cells:
                set_cell_shading(cell, 'FFFFFF')
    else:
        document.add_paragraph("無最終判定記錄。")


    # 8. 執行結果
    document.add_heading("執行結果", level=2)
    if case.results.exists():
        table_result = document.add_table(rows=1, cols=5)
        table_result.autofit = True
        table_result.style = 'Table Grid' # Add this line
        hdr_cells_result = table_result.rows[0].cells
        hdr_cells_result[0].text = "應買止日"
        hdr_cells_result[1].text = "執行結果"
        hdr_cells_result[2].text = "搶標拍別"
        hdr_cells_result[3].text = "搶標金額"
        hdr_cells_result[4].text = "標的編號"
        # Make headers bold and set background to light orange
        for cell in hdr_cells_result:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            set_cell_shading(cell, 'FFDAB9')

        # Define specific column widths for Result table
        specific_column_widths_cm_result = [
            5.08,  # 應買止日
            5.08,  # 執行結果
            5.08,  # 搶標拍別
            5.08,  # 搶標金額
            5.08,  # 標的編號
        ]

        # Set column widths for each cell (column-wise iteration)
        for col_idx, width_cm in enumerate(specific_column_widths_cm_result):
            for cell in table_result.columns[col_idx].cells:
                cell.width = Cm(width_cm)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for result in case.results.all():
            row_cells = table_result.add_row().cells
            row_cells[0].text = result.stopBuyDate.strftime("%Y-%m-%d") if result.stopBuyDate else '—'
            row_cells[1].text = result.actionResult or '—'
            row_cells[2].text = result.bidAuctionTime or '—'
            row_cells[3].text = str(result.bidMoney) or '—'
            row_cells[4].text = result.objectNumber or '—'
            # Set data row white background
            for cell in row_cells:
                set_cell_shading(cell, 'FFFFFF')
    else:
        document.add_paragraph("無結果記錄。")


    # 9. 比價建物 (Object Builds) - Complex Section
    document.add_heading("比價建物", level=2)
    if case.objectbuilds.exists():
        table_objectbuild = document.add_table(rows=1, cols=10) # Changed from 11 to 10 columns
        table_objectbuild.autofit = True
        table_objectbuild.style = 'Table Grid' # Add this line
        hdr_cells_objectbuild = table_objectbuild.rows[0].cells
        hdr_cells_objectbuild[0].text = "類型"
        hdr_cells_objectbuild[1].text = "地址"
        hdr_cells_objectbuild[2].text = "屋齡(年)"
        hdr_cells_objectbuild[3].text = "成交日期"
        hdr_cells_objectbuild[4].text = "樓高"
        hdr_cells_objectbuild[5].text = "總價"
        hdr_cells_objectbuild[6].text = "建坪(坪)"
        hdr_cells_objectbuild[7].text = "增坪(坪)"
        hdr_cells_objectbuild[8].text = "單價"
        hdr_cells_objectbuild[9].text = "計算(加成)"
        # Make headers bold and set background to light orange
        for cell in hdr_cells_objectbuild:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            set_cell_shading(cell, 'FFDAB9')

        # Define specific column widths for Result table
        specific_column_widths_cm_result = [
            2.25,
            4.93,   
            2.12,  
            2.68,  
            1.69,  
            2.21, 
            2.45,  
            2.45, 
            2.45,  
            2.45, 
        ]

        # Set column widths for each cell (column-wise iteration)
        for col_idx, width_cm in enumerate(specific_column_widths_cm_result):
            for cell in table_objectbuild.columns[col_idx].cells:
                cell.width = Cm(width_cm)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for objectbuild in case.objectbuilds.all():
            row_cells = table_objectbuild.add_row().cells
            row_cells[0].text = objectbuild.type or '—'
            
            address_text = objectbuild.address or '—'
            if objectbuild.url and objectbuild.url != '—':
                # Clear existing content and add hyperlink to the address cell
                row_cells[1].text = '' # Clear existing text
                add_hyperlink(row_cells[1].paragraphs[0], address_text, objectbuild.url)
            else:
                row_cells[1].text = address_text

            row_cells[2].text = str(objectbuild.houseAge) or '—'
            row_cells[3].text = str(objectbuild.transactionDate) or '—'
            row_cells[4].text = str(objectbuild.floorHeight) or '—'
            row_cells[5].text = str(objectbuild.totalPrice) or '—'
            row_cells[6].text = str(objectbuild.buildArea) or '—'
            row_cells[7].text = str(objectbuild.subBuildArea) or '—'
            row_cells[8].text = str(objectbuild.unitPrice) or '—'
            row_cells[9].text = str(objectbuild.calculate) or '—'
            # Set data row white background
            for cell in row_cells:
                set_cell_shading(cell, 'FFFFFF')
    else:
        document.add_paragraph("無結果記錄。")


    # 10. 公文 (Official Documents)
    document.add_heading("公文", level=2)
    if case.officialdocuments.exists():
        table_od = document.add_table(rows=1, cols=5) # Changed to 5 columns
        table_od.autofit = True
        table_od.style = 'Table Grid' # Add this line
        # Add headers for Official Document table
        hdr_cells_od = table_od.rows[0].cells
        hdr_cells_od[0].text = "案別"
        hdr_cells_od[1].text = "股別"
        hdr_cells_od[2].text = "案號"
        hdr_cells_od[3].text = "電話"
        hdr_cells_od[4].text = "分機"
        # Make headers bold and set background to light orange
        for cell in hdr_cells_od:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
            set_cell_shading(cell, 'FFDAB9')

        # Define specific column widths for Official Document table
        specific_column_widths_cm_od = [
            5.09, # 案別
            5.09, # 股別
            5.09, # 案號
            5.09, # 電話
            5.09, # 分機
        ]

        # Set column widths for each cell (column-wise iteration)
        for col_idx, width_cm in enumerate(specific_column_widths_cm_od):
            for cell in table_od.columns[col_idx].cells:
                cell.width = Cm(width_cm)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for od in case.officialdocuments.all():
            row_cells = table_od.add_row().cells
            row_cells[0].text = od.type or '—' # Corrected from od.documentType
            row_cells[1].text = od.stock or '—' # Corrected from od.branchNumber
            row_cells[2].text = od.docNumber or '—' # Corrected from od.caseNumber
            row_cells[3].text = od.tel or '—' # Corrected from od.phoneNumber
            row_cells[4].text = od.ext or '—' # Corrected from od.extensionNumber
            # Set data row white background
            for cell in row_cells:
                set_cell_shading(cell, 'FFFFFF')
    else:
        document.add_paragraph("無公文記錄。")

    # Compare Building (比價建物) Section
    document.add_heading("比價建物", level=2)
    if case.objectbuilds.exists(): # Corrected from case.compares.exists()
        document.add_paragraph("無比價建物記錄。") # Temporarily bypass table creation
    else:
        document.add_paragraph("無比價建物記錄。")


    f = BytesIO()
    document.save(f)
    f.seek(0)

    response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    # Constructing the Content-Disposition header with only filename* for better UTF-8 handling
    encoded_filename_utf8 = quote(base_filename, encoding='utf-8')
    response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename_utf8}"
    return response

